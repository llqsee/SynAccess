from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "reviewer_response_anomaly_algorithm.docx"


# ---------------------------------------------------------------------------
# Response text. Every quantitative statement below is grounded in
# backend/services/anomaly_detection_service.py (the authoritative
# implementation) rather than restated from memory.
# ---------------------------------------------------------------------------

REBUTTAL = (
    "We thank the reviewer for this valuable comment. We agree that the original "
    "manuscript under-specified the anomaly-detection procedure. In the revision we now "
    "present (i) the complete algorithm as numbered steps, (ii) its statistical "
    "foundation and the derivation of the per-cell test, (iii) explicit criteria for "
    "selecting each parameter, and (iv) a formal time- and space-complexity analysis. "
    "The method is deliberately simple by design: rather than a learned density model, it "
    "treats the labelling of points as real vs. synthetic within a spatial grid cell as a "
    "Bernoulli process and tests, cell by cell, whether the local mixing proportion departs "
    "from the global proportion, with multiplicity controlled by a false-discovery-rate "
    "procedure. This makes every reported anomaly a statistically interpretable, "
    "multiplicity-corrected hypothesis-test outcome rather than an opaque score. The "
    "following material has been added to the Methods section and a new appendix."
)


# --- Section 1: complete algorithm steps ----------------------------------

ALGORITHM_INTRO = (
    "Let R = {r_1, ..., r_{m}} be the real points and S = {s_1, ..., s_{k}} the synthetic "
    "points, each embedded in the same 2-D space (the UMAP/t-SNE embedding). Write "
    "N = m + k for the total number of points. The procedure is:"
)

ALGORITHM_STEPS = [
    "Step 1 - Grid construction. Pool the real and synthetic points and form one-"
    "dimensional histograms of the x- and y-coordinates independently, using B_x and B_y "
    "equal-width bins respectively (numpy.histogram). This yields bin edges "
    "e^x_0 < e^x_1 < ... < e^x_{B_x} and e^y_0 < ... < e^y_{B_y}. The outer edges are "
    "padded outward by 1% of each coordinate's range so that the extreme points fall "
    "strictly inside the grid. The Cartesian product of the x- and y-bins defines "
    "G = B_x * B_y rectangular cells C_{ij}.",

    "Step 2 - Global proportion. Compute the global real proportion "
    "p_0 = m / (m + k). This is the expected fraction of real points in any region under "
    "the null hypothesis that real and synthetic points are spatially interchangeable.",

    "Step 3 - Per-cell counts. For each cell C_{ij}, count the real points r_{ij} and "
    "synthetic points n^s_{ij} that fall inside it, and set the cell total "
    "n_{ij} = r_{ij} + n^s_{ij}. Cells with n_{ij} < tau (tau = 5 by default) are skipped "
    "to avoid testing on too little evidence.",

    "Step 4 - Two one-sided exact binomial tests. For every retained cell, treat r_{ij} as "
    "a draw from Binomial(n_{ij}, p_0) and compute two one-sided exact binomial p-values: "
    "a 'greater' test for real over-population (local proportion above p_0) and a 'less' "
    "test for synthetic over-population (local proportion below p_0). Each cell is routed "
    "to the positive family if its observed proportion p_hat_{ij} = r_{ij}/n_{ij} exceeds "
    "p_0, and to the negative family if it is below p_0.",

    "Step 5 - Family-wise FDR correction. The two families (positive = real over-"
    "population, negative = synthetic over-population) are corrected separately with the "
    "Benjamini-Hochberg procedure at level alpha (fdr_alpha = 0.05 by default). A cell is "
    "declared a significant anomaly only if its BH-adjusted p-value is at or below alpha.",

    "Step 6 - Labelling and visualisation. Significant real-over-populated cells are "
    "coloured red, significant synthetic-over-populated cells blue, all others neutral. "
    "Every point inherits the anomaly status of the cell it falls in, giving point-level "
    "real/synthetic anomaly sets and the aggregate anomaly rates reported to the user.",
]


# --- Section 2: statistical / mathematical basis ---------------------------

DERIVATION_INTRO = (
    "Statistical model and derivation of the per-cell test."
)

DERIVATION_PARAS = [
    "Null hypothesis. Fix a cell C_{ij} containing n_{ij} points. Under the null "
    "hypothesis H0 that real and synthetic points are drawn from the same spatial "
    "distribution, the label (real / synthetic) of each point in the cell is an "
    "independent Bernoulli trial with success probability p_0 = m/(m+k), independent of "
    "location. Consequently the number of real points in the cell is the sum of n_{ij} "
    "independent Bernoulli(p_0) variables, i.e.",

    "    r_{ij} | n_{ij}  ~  Binomial(n_{ij}, p_0)   under H0.",

    "This is exactly the conditional distribution one obtains by conditioning the two "
    "independent Poisson/Multinomial point counts on their cell total, which is why the "
    "binomial - rather than a normal approximation - is the natural exact reference "
    "distribution.",

    "Alternative hypotheses and p-values. A region where the generator under-samples real "
    "structure shows an excess of real points (p_hat_{ij} > p_0); a region where it over-"
    "generates shows a deficit (p_hat_{ij} < p_0). These are one-sided departures, so we "
    "use two one-sided exact binomial tests. For the real-over-population (greater) "
    "alternative the exact p-value is the upper tail",

    "    P(X >= r_{ij}) = sum_{x=r_{ij}}^{n_{ij}} C(n_{ij}, x) p_0^x (1-p_0)^{n_{ij}-x},",

    "and symmetrically for the synthetic-over-population (less) alternative the lower tail "
    "P(X <= r_{ij}) is used. The exact test is chosen over the normal/score approximation "
    "because cell counts are small and p_0 can be far from 0.5, where the approximation is "
    "unreliable; the exact tail probability is valid for any n_{ij} >= 1.",

    "Effect size. Significance is reported together with the signed effect size "
    "p_hat_{ij} - p_0, so that statistically significant but practically negligible "
    "deviations can be distinguished from substantively large ones.",

    "Multiplicity control. Because one test is performed per retained cell, G' tests are "
    "conducted and the family-wise error would be severely inflated if raw p-values were "
    "thresholded at alpha. We control the false discovery rate with the Benjamini-Hochberg "
    "(BH) procedure: order the p-values p_(1) <= ... <= p_(G'), find the largest l with "
    "p_(l) <= (l / G') * alpha, and reject all hypotheses with p-value <= p_(l). Under "
    "independence (and more generally positive regression dependence) BH guarantees "
    "E[FDR] <= alpha. The positive and negative families are corrected independently "
    "because they answer two distinct one-sided questions and pooling them would conflate "
    "the two directions of deviation.",

    "Interpretation. The output is therefore a set of grid cells, each carrying an exact "
    "p-value, a BH-adjusted p-value, a signed effect size, and a binary significance flag. "
    "An 'anomaly' is precisely a cell in which the local real/synthetic mixing proportion "
    "differs from the global proportion by more than sampling noise after FDR control - a "
    "fully reproducible and interpretable definition.",
]


# --- Section 3: parameter selection criteria -------------------------------

PARAMETER_INTRO = (
    "The method has four parameters. Their defaults, admissible ranges, and the criteria "
    "governing their choice are as follows."
)

PARAMETER_POINTS = [
    "Number of bins per axis B_x, B_y (default 20, range 5-100). This is the resolution / "
    "power trade-off. The grid has G = B_x * B_y cells, so the expected count per cell is "
    "roughly N / G. To keep most cells above the minimum-count threshold tau, the bin "
    "count should satisfy B_x * B_y <~ N / tau; equivalently, for a roughly square grid, "
    "B per axis <~ sqrt(N / tau). Too few bins blur localised discrepancies; too many "
    "leave most cells below tau (untested) and reduce per-test power. The default of 20 "
    "per axis is appropriate for the few-thousand-point embeddings typical of this "
    "application.",

    "Minimum cell count tau (default 5). Cells with fewer than tau points are not tested. "
    "Five is the conventional lower bound for a meaningful proportion test; because the "
    "test is exact, smaller cells are not invalid, merely under-powered, so tau is a power "
    "floor rather than a validity requirement. Raising tau yields fewer, more reliable "
    "tests; lowering it increases spatial resolution at the cost of power.",

    "FDR level alpha (fdr_alpha, default 0.05, range 0.001-0.5). This is the tolerated "
    "expected proportion of false discoveries among the flagged cells. Lower alpha is more "
    "conservative (fewer false anomalies, lower sensitivity); higher alpha is more "
    "sensitive. 0.05 is the standard default and is recommended unless the downstream use "
    "is screening (higher alpha) or confirmatory (lower alpha).",

    "Edge padding (1% of range). A purely numerical safeguard that guarantees the extreme "
    "points are enclosed by the outer bins; it has no effect on the statistics and is not "
    "intended as a tuning knob.",
]


# --- Section 4: complexity analysis ----------------------------------------

COMPLEXITY_INTRO = (
    "Complexity analysis. Let N = m + k be the number of points and G = B_x * B_y the "
    "number of grid cells."
)

COMPLEXITY_POINTS = [
    "Grid construction. Building the two 1-D histograms is O(N) time and O(B_x + B_y) "
    "space for the edges.",

    "Per-cell counting. The reference implementation counts each cell by scanning the "
    "point arrays, costing O(N) per cell and O(N * G) overall. This is the dominant term. "
    "It can be reduced to O(N + G) by assigning every point to its cell once via vectorised "
    "digitisation (numpy.digitize) and accumulating a 2-D count matrix - an optimisation we "
    "note in the revision; point-level cell assignment already uses this O(N) digitisation.",

    "Hypothesis testing. One exact binomial test per retained cell; each tail sum is "
    "O(n_{ij}) in the worst case, bounded overall by O(N) work across all cells, plus the "
    "fixed cost of at most O(G) tests.",

    "FDR correction. Sorting the p-values within each family dominates this stage at "
    "O(G log G).",

    "Totals. Time is O(N * G) for the reference implementation, or O(N + G log G) with the "
    "vectorised-counting optimisation. Space is O(N + G): O(N) for the point-to-cell "
    "assignments and O(G) for per-cell counts, p-values and adjusted p-values. The method "
    "is therefore linear in the data size (up to the grid factor) and trivially "
    "parallelisable across cells; an optional GPU path is provided for the histogram and "
    "digitisation steps on large inputs.",
]


MANUSCRIPT_CHANGES = [
    "Added a numbered algorithm box (Steps 1-6) to the Methods section.",
    "Added a subsection deriving the binomial null model, the two one-sided exact tests, "
    "and the Benjamini-Hochberg FDR control, with the alternative hypotheses made explicit.",
    "Added a parameter table giving defaults, ranges, and selection criteria for B_x, B_y, "
    "tau and alpha.",
    "Added a complexity paragraph stating O(N + G log G) achievable time and O(N + G) space, "
    "and noting the reference O(N * G) counting loop and its vectorised optimisation.",
]


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paras(document: Document, paras) -> None:
    for p in paras:
        document.add_paragraph(p)


def add_bullets(document: Document, points) -> None:
    for p in points:
        document.add_paragraph(p, style="List Bullet")


def add_numbered(document: Document, points) -> None:
    for p in points:
        document.add_paragraph(p, style="List Number")


def main() -> None:
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(
        "Reviewer Response: Mathematical Derivation, Parameter Selection, "
        "and Complexity of the Histogram-Based Anomaly-Detection Algorithm"
    )
    title_run.bold = True
    title_run.font.size = Pt(15)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("Prepared from the current SynAccess implementation")

    add_heading(document, "Response", level=1)
    document.add_paragraph(REBUTTAL)

    add_heading(document, "1. Complete Algorithm Steps", level=1)
    document.add_paragraph(ALGORITHM_INTRO)
    add_numbered(document, ALGORITHM_STEPS)

    add_heading(document, "2. Statistical and Mathematical Basis", level=1)
    add_paras(document, DERIVATION_PARAS)

    add_heading(document, "3. Parameter Selection Criteria", level=1)
    document.add_paragraph(PARAMETER_INTRO)
    add_bullets(document, PARAMETER_POINTS)

    add_heading(document, "4. Complexity Analysis", level=1)
    document.add_paragraph(COMPLEXITY_INTRO)
    add_bullets(document, COMPLEXITY_POINTS)

    add_heading(document, "5. Summary of Manuscript Changes", level=1)
    add_bullets(document, MANUSCRIPT_CHANGES)

    output_path = OUTPUT
    try:
        document.save(output_path)
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = ROOT / f"reviewer_response_anomaly_algorithm_{timestamp}.docx"
        document.save(output_path)

    print(output_path)


if __name__ == "__main__":
    main()
