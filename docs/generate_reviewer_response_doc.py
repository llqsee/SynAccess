from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "reviewer_response_ai_grounding.docx"


REBUTTAL = (
    "We appreciate this concern and agree that the prompt engineering and grounding strategy "
    "should be described more explicitly. In the current implementation, the LLM does not "
    "generate statistical test results itself; rather, it receives the structured validation-results "
    "JSON produced by the backend statistical pipeline, which deterministically computes metrics "
    "such as Kolmogorov-Smirnov, chi-square, and Welch's t-test p-values, with Benjamini-Hochberg "
    "false-discovery-rate correction applied where appropriate. The AI analysis layer then prompts "
    "the model to base its interpretation only on these provided results and to reference the "
    "statistical tests and p-values explicitly. That said, the current version relies primarily on "
    "prompt-level constraints and does not yet include a fully programmatic claim-verification layer "
    "to ensure that every narrative statement is traceable to a computed value before it is returned "
    "to the user. We will therefore revise the manuscript to include the full prompt in an appendix, "
    "clarify the exact grounding mechanism, and explicitly state this limitation: the LLM is intended "
    "as a narrative interface over computed evidence, not as an independent source of statistical inference."
)


GROUNDING_POINTS = [
    "The statistical values are computed in the backend validation pipeline, not by the LLM.",
    "The narrative model receives the validation-results JSON as evidence context.",
    "The prompt instructs the model not to hallucinate and to cite p-values and tests.",
    "No claim-level checker currently verifies that each generated statement maps to a specific field in the structured output.",
    "The present safeguard is therefore prompt-based grounding, not hard post-generation validation.",
]


NARRATIVE_SYSTEM_PROMPT = (
    "You are a senior Data Quality Expert. Analyze the validation results and give clear, useful recommendations."
)


NARRATIVE_USER_PROMPT_TEMPLATE = """You are a senior Data Quality Expert.

The following are validation results comparing real and synthetic data:

{{json.dumps(validation_results, indent=2)}}

Please provide a clear analysis that includes:
SECTION_TITLE: Summary
In the Summary section, include one bullet point titled 'AI Generated Score:' with these six scores (0-100): Overall Similarity Score, Multi-variable Score, Anomaly Burden Score, Bivariate Score, Structural Validity Score, and Univariate Similarity Score.
SECTION_TITLE: Key Findings (50 words maximum)
SECTION_TITLE: Statistical Quality (50 words maximum)
SECTION_TITLE: Practical Usefulness (50 words maximum)
SECTION_TITLE: Critical Issues (if any)
SECTION_TITLE: Variable Range and Type Checks: identify synthetic variables with unreasonable value ranges and flag any data type mismatches between real and synthetic data, highlighting the variables.
SECTION_TITLE: Real Data Quality by Variable: assess each real-data variable for plausible ranges and show the variables in real data that have unreasonable values based on domain realities.
SECTION_TITLE: Actionable Recommendations (3-5 recommendations, (50 words maximum))
SECTION_TITLE: Risk Level (LOW, MEDIUM, HIGH) with justification (50 words maximum))

IMPORTANT: Start each section with the marker 'SECTION_TITLE: [Title]' on a new line and then provide the content. Do NOT use markdown bold stars (**) for these titles. This is critical for PDF generation.
IMPORTANT: Under each section title, provide 3-5 bullet points. Each bullet must begin with '- ' and contain no more than 12 words. Do not write long paragraphs.
Dont Hallucinate any data, only base your analysis on the validation results and the data itself. Always justify your findings and recommendations using the data itself. Make sure to include p-values and reference the statistical tests used.
"""


SIMILARITY_SYSTEM_PROMPT = (
    "You are a senior data quality statistician. Return only valid JSON with numeric scores from 0 to 100."
)


SIMILARITY_USER_PROMPT_TEMPLATE = """Given these dataset similarity metrics, provide AI-estimated counterparts.
Keep values on a 0-100 scale and consistent with the evidence.

{{json.dumps({\"computed_scores\": computed_scores or {}, \"context\": context or {}}, indent=2)}}

Output ONLY JSON with this exact schema:
{
  \"overallScore\": number,
  \"structuralScore\": number,
  \"univariateScore\": number,
  \"bivariateScore\": number,
  \"anomalyScore\": number,
  \"multiVariableScore\": number,
  \"confidence\": \"LOW|MEDIUM|HIGH\",
  \"rationale\": \"short explanation\"
}
"""


def add_code_block(document: Document, text: str) -> None:
    for line in text.splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)


def main() -> None:
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("Reviewer Response: AI Prompt Transparency and Grounding")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("Prepared from the current SynAccess implementation")

    document.add_heading("Response", level=1)
    document.add_paragraph(REBUTTAL)

    document.add_heading("Grounding Mechanism Summary", level=1)
    for point in GROUNDING_POINTS:
        document.add_paragraph(point, style="List Bullet")

    document.add_heading("Appendix A. Narrative Analysis Prompt", level=1)
    document.add_paragraph("System prompt used for the narrative AI analysis call:")
    add_code_block(document, NARRATIVE_SYSTEM_PROMPT)
    document.add_paragraph(
        "User prompt template constructed in the service (with the runtime validation JSON inserted in place of the placeholder):"
    )
    add_code_block(document, NARRATIVE_USER_PROMPT_TEMPLATE)

    document.add_heading("Appendix B. Similarity Scoring Prompt", level=1)
    document.add_paragraph("System prompt used for the similarity-score call:")
    add_code_block(document, SIMILARITY_SYSTEM_PROMPT)
    document.add_paragraph("User prompt template for the similarity-score request:")
    add_code_block(document, SIMILARITY_USER_PROMPT_TEMPLATE)

    document.add_heading("Appendix C. Interpretation Note", level=1)
    document.add_paragraph(
        "The appendix reflects the prompts as implemented in the backend service. It should be read alongside the description of the deterministic statistical pipeline, because the p-values and FDR-corrected results are computed before the LLM is called."
    )

    output_path = OUTPUT
    try:
        document.save(output_path)
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = ROOT / f"reviewer_response_ai_grounding_{timestamp}.docx"
        document.save(output_path)

    print(output_path)


if __name__ == "__main__":
    main()